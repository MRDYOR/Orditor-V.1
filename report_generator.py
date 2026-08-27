"""
Orditor AIX — pure-Python docx report generator.

Reimplements the same report shape as the original JS/docx-npm generator, but
using python-docx, so the whole app runs in one runtime (important for a
serverless deploy where shelling out to Node isn't reliably available).

report_bytes(findings_dict, meta_dict) -> bytes  (the .docx file content)
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SEV_COLOR = {
    "Critical": RGBColor(0x8B, 0x1E, 0x1E), "High": RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xB9, 0x77, 0x0E), "Low": RGBColor(0x7D, 0x66, 0x08),
    "Informational": RGBColor(0x55, 0x55, 0x55),
}
SEV_FILL = {
    "Critical": "F5D5D5", "High": "F8E0DA", "Medium": "FBEBCE",
    "Low": "F5F1D8", "Informational": "ECECEC",
}
SURFACE_COLOR = {
    "Agent": RGBColor(0x6B, 0x4F, 0xBB), "Application": RGBColor(0x1F, 0x6F, 0xB2),
    "Blockchain": RGBColor(0x1E, 0x7A, 0x4C), "Integration": RGBColor(0xB2, 0x70, 0x1F),
}
SEV_ORDER = ["Critical", "High", "Medium", "Low", "Informational"]
SURFACE_ORDER = ["Agent", "Integration", "Application", "Blockchain"]


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _para(doc, text, italic=False, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def _risk_rating(counts):
    if counts.get("Critical"):
        return "CRITICAL RISK", SEV_COLOR["Critical"]
    if counts.get("High"):
        return "HIGH RISK", SEV_COLOR["High"]
    if counts.get("Medium"):
        return "MEDIUM RISK", SEV_COLOR["Medium"]
    if counts.get("Low"):
        return "LOW RISK", SEV_COLOR["Low"]
    return "NO FINDINGS", RGBColor(0x2E, 0x7D, 0x32)


def report_bytes(data: dict, meta: dict) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(10)

    # ---- Cover ----
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ORDITOR")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x8B, 0x7C, 0xF6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AIX \u2014 AI / Web3 Security Audit Framework")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Unified Security Report")
    run.bold = True
    run.font.size = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("projectName", "Unnamed Project"))
    run.font.size = Pt(16)

    for _ in range(2):
        doc.add_paragraph()

    for label, val in [
        ("Files scanned", len(data.get("targets_scanned", []))),
        ("Chain", meta.get("chain")),
        ("Report date", meta.get("auditDate", "")),
        ("Report version", meta.get("reportVersion", "1.0")),
    ]:
        if val:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}: {val}")
            run.font.size = Pt(11)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Automated multi-surface static-analysis review \u2014 covers Application, Agent, "
        "Blockchain, and Integration (AI\u2192blockchain boundary) surfaces. Not a substitute "
        "for a full manual audit. See Methodology & Limitations."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # ---- Executive summary ----
    counts = data.get("by_severity", {})
    total = data.get("finding_count", 0)
    by_surface = data.get("by_surface", {})

    _heading(doc, "Executive Summary", 1)
    _para(doc, (
        f"Orditor performed an automated, multi-surface static-analysis review of "
        f"{meta.get('projectName', 'the submitted project')}. {total} finding(s) were "
        f"identified across {len(data.get('targets_scanned', []))} file(s), spanning "
        f"{len(by_surface)} of the framework's four code-level surfaces (Agent, Application, "
        f"Blockchain, Integration \u2014 the remaining two, Economics and Business Logic, "
        f"require human/manual review; see Limitations)."
    ))
    label, color = _risk_rating(counts)
    p = doc.add_paragraph()
    p.add_run("Overall rating: ").bold = True
    run = p.add_run(label)
    run.bold = True
    run.font.color.rgb = color

    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    _cell_text(hdr[0], "Severity", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _shade_cell(hdr[0], "222222")
    _cell_text(hdr[1], "Count", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _shade_cell(hdr[1], "222222")
    for sev in SEV_ORDER:
        row = t.add_row().cells
        _cell_text(row[0], sev, bold=True, color=SEV_COLOR[sev])
        _shade_cell(row[0], SEV_FILL[sev])
        _cell_text(row[1], str(counts.get(sev, 0)))
        _shade_cell(row[1], SEV_FILL[sev])

    doc.add_paragraph()
    t2 = doc.add_table(rows=1, cols=2)
    hdr = t2.rows[0].cells
    _cell_text(hdr[0], "Surface", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _shade_cell(hdr[0], "222222")
    _cell_text(hdr[1], "Findings", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _shade_cell(hdr[1], "222222")
    for s in SURFACE_ORDER:
        if by_surface.get(s):
            row = t2.add_row().cells
            _cell_text(row[0], s, bold=True, color=SURFACE_COLOR.get(s))
            _cell_text(row[1], str(by_surface[s]))

    doc.add_paragraph()
    _para(doc, (
        "This rating reflects the highest-severity class of finding present. It is not a "
        "probability estimate and does not account for mitigating controls (e.g. a multisig "
        "timelock, an infrastructure-level rate limiter) unless those are reflected in the "
        "code itself."
    ))

    # ---- Attack chain ----
    _heading(doc, "Attack Chain Analysis", 1)
    notes = data.get("attack_chain_notes", [])
    if notes:
        _para(doc, (
            "Per the framework's Phase 7 principle: individual components may each look "
            "acceptable in isolation. The combinations below are synthesized because "
            "Critical/High findings appear together across surfaces in ways that may compose "
            "into a single exploitable path \u2014 each is worth a manual chained-exploit "
            "review, not just an individual fix."
        ), italic=True)
        for n in notes:
            _bullet(doc, n)
    else:
        _para(doc, (
            "No Critical/High findings spanning multiple surfaces were detected, so no "
            "chained-exploit pattern was synthesized. This does not rule out chains between "
            "Medium/Low findings or between findings and factors outside static analysis's "
            "reach (business logic, economic incentives) \u2014 see Limitations."
        ), italic=True)

    # ---- Methodology ----
    _heading(doc, "Methodology & Scope", 1)
    _para(doc, "Approach: automated static analysis across four engines, one per surface the AIX framework defines at the code level:")
    for t_ in [
        "Blockchain \u2014 Solidity source: reentrancy, access control, dangerous operations, arithmetic, centralization, code hygiene.",
        "Application \u2014 Python API/backend source: auth gaps, secrets management, injection, CORS, configuration, error handling, data handling, rate limiting.",
        "Agent \u2014 Python agent code: containment (eval/exec on model output, tool-dispatch shell access, unbounded loops, outbound network access) and payment verification (signature, replay/nonce, amount/recipient, expiry).",
        "Integration \u2014 the AI\u2192blockchain boundary specifically: does an AI/agent-decided value (recipient, amount, contract, function) reach transaction construction/signing without independent deterministic validation.",
    ]:
        _bullet(doc, t_)
    p = doc.add_paragraph()
    p.add_run("Out of scope for this pass (per the framework, these require human/manual reasoning, not static analysis):").bold = True
    for t_ in [
        "Phase 1 Threat Modeling \u2014 who can attack what, through which path \u2014 requires architectural context this tool doesn't have.",
        "Phase 6 Business Logic \u2014 legitimate functionality used in an unintended way (double-claiming, referral abuse, timing exploits) doesn't match a fixed code pattern.",
        "Economic Security \u2014 price manipulation, oracle attacks, flash-loan attacks, MEV, incentive design.",
        "Model manipulation \u2014 prompt injection / context poisoning susceptibility is a property of the model and its inputs, not the surrounding code.",
        "Anything requiring live execution: dynamic testing, fuzzing, invariant testing, transaction simulation.",
    ]:
        _bullet(doc, t_)
    _para(doc, meta.get("scopeNote", "Scope was limited to the files provided."), italic=True)

    _heading(doc, "Finding Classification (Phase 9)", 2)
    _para(doc, (
        "Each finding carries: Severity (Critical/High/Medium/Low/Informational), Confidence "
        "(Confirmed \u2014 a directly verifiable fact about the source text, e.g. a literal "
        "hardcoded key; or Probable \u2014 a heuristic pattern match that should be read, not "
        "just counted), Impact (Funds/Integrity/Availability/Privacy/Privilege/Reputation), and "
        "Surface (which of the four code-level domains above it belongs to)."
    ))

    _heading(doc, "Limitations", 2)
    for t_ in [
        "This is heuristic, pattern-based static analysis across all four engines \u2014 not a full taint/dataflow/symbolic-execution tool. Every engine will produce false positives and can miss issues that don't match a known pattern.",
        "No live requests, no transaction simulation, no dynamic testing was performed.",
        "The AI\u2192blockchain boundary check (Integration surface) confirms whether something validation-shaped exists between an AI decision and a transaction \u2014 not whether that validation is correct or complete. Read every Integration finding manually.",
        "Attack-chain notes are synthesized from co-occurrence of severe findings across surfaces \u2014 they flag combinations worth a manual look, not confirmed exploit chains.",
        "This report is not a substitute for a full manual audit, and particularly not for anything with real funds, real user data, or real agent-initiated transactions in production.",
    ]:
        _bullet(doc, t_)

    doc.add_page_break()

    # ---- Findings index ----
    _heading(doc, "Findings Index", 1)
    findings = data.get("findings", [])
    t3 = doc.add_table(rows=1, cols=5)
    hdr = t3.rows[0].cells
    for i, label in enumerate(["ID", "Title", "Severity", "Surface", "File"]):
        _cell_text(hdr[i], label, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr[i], "222222")
    for f in findings:
        row = t3.add_row().cells
        _cell_text(row[0], f.get("id", ""))
        _cell_text(row[1], f.get("title", ""))
        _cell_text(row[2], f.get("severity", ""), bold=True, color=SEV_COLOR.get(f.get("severity")))
        _shade_cell(row[2], SEV_FILL.get(f.get("severity"), "FFFFFF"))
        _cell_text(row[3], f.get("surface", ""), bold=True, color=SURFACE_COLOR.get(f.get("surface")))
        _cell_text(row[4], (f.get("file") or "").split("/")[-1])

    doc.add_page_break()

    # ---- Detailed findings ----
    _heading(doc, "Detailed Findings", 1)
    if not findings:
        _para(doc, "No findings were produced by this pass.")
    for f in findings:
        _heading(doc, f"{f.get('id','')} \u2014 {f.get('title','')}", 3)
        loc = f"Line {f['line']}" if f.get("line") else "File-wide"
        fn = f" \u2014 function `{f['function']}`" if f.get("function") else ""
        info = [
            ("ID", f.get("id", "")), ("Severity", f.get("severity", "")),
            ("Confidence", f.get("confidence", "")), ("Impact", f.get("impact", "")),
            ("Surface", f.get("surface", "")), ("Category", f.get("category", "")),
            ("File", (f.get("file") or "").split("/")[-1]), ("Location", f"{loc}{fn}"),
        ]
        it = doc.add_table(rows=0, cols=2)
        for label, val in info:
            row = it.add_row().cells
            _cell_text(row[0], label, bold=True)
            _shade_cell(row[0], "F2F2F2")
            color = SEV_COLOR.get(val) if label == "Severity" else (SURFACE_COLOR.get(val) if label == "Surface" else None)
            _cell_text(row[1], str(val), bold=(label in ("Severity", "Surface")), color=color)
            if label == "Severity":
                _shade_cell(row[1], SEV_FILL.get(val, "FFFFFF"))
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(f.get("description", ""))
        if f.get("snippet"):
            p = doc.add_paragraph()
            run = p.add_run(f["snippet"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        p = doc.add_paragraph()
        p.add_run("Recommendation: ").bold = True
        p.add_run(f.get("recommendation", ""))
        doc.add_paragraph()

    doc.add_page_break()

    # ---- Remediation summary ----
    _heading(doc, "Remediation Summary", 1)
    _para(doc, (
        "Grouped by surface, highest severity first within each. Fix root causes, not just "
        "symptoms \u2014 several findings within a surface often share one underlying gap."
    ))
    by_surface_findings = {}
    for f in findings:
        if f.get("severity") == "Informational":
            continue
        by_surface_findings.setdefault(f.get("surface"), []).append(f)
    for s in SURFACE_ORDER:
        if s not in by_surface_findings:
            continue
        _heading(doc, s, 2)
        group = sorted(by_surface_findings[s], key=lambda f: SEV_ORDER.index(f.get("severity", "Low")))
        for f in group:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{f.get('severity')}] ")
            run.bold = True
            run.font.color.rgb = SEV_COLOR.get(f.get("severity"))
            p.add_run(f"{f.get('title')}: ")
            run2 = p.add_run(f.get("recommendation", ""))
            run2.italic = True

    # ---- Disclaimer ----
    _heading(doc, "Disclaimer", 1)
    _para(doc, (
        "This report reflects an automated review performed by Orditor and is provided for "
        "informational purposes only. It does not constitute a certification that the "
        "reviewed project is secure, and it is not financial, investment, or legal advice. "
        "Findings marked at any severity level should be independently verified and, for any "
        "project handling real funds, real user data, or real agent-initiated transactions, "
        "reviewed by a qualified human security team before production deployment. Orditor "
        "and its operator accept no liability for losses arising from reliance on this report."
    ))
    _para(doc, (
        "Business logic, economic security, threat modeling, and model-manipulation "
        "resistance are explicitly out of scope for this automated pass \u2014 see "
        "Methodology & Limitations."
    ), italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
